from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt
import io


def add_horizontal_line(paragraph):
    """
    Adds a horizontal line by setting the paragraph bottom border.
    """
    p = paragraph._p  # Access the underlying lxml paragraph element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert(0, pBdr)
    
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')  # Correctly use qn for namespaced attribute
    bottom_bdr.set(qn('w:sz'), '4')  # Size of border
    bottom_bdr.set(qn('w:space'), '1')  # Spacing of border
    bottom_bdr.set(qn('w:color'), 'auto')  # Color of border
    pBdr.append(bottom_bdr)

def qn(val):
    """
    WordprocessingML utility function for namespaced element creation.
    """
    return '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}%s' % val


def generate_docx_for_brief(brief_data):
    doc = Document()
    
    

    doc = generate_cover_page(doc, brief_data)
    doc.add_page_break()

    
    # Questions Presented
    if 'questions_presented' in brief_data and brief_data['questions_presented']:
        add_custom_heading(doc, 'Questions Presented', level=1, font_size=18)
        for i, each in enumerate(brief_data['questions_presented'], start=1):
            add_custom_paragraph(doc, f"{i}. {each}", font_size=16)

    # Table of Authorities
    if 'table_of_authorities' in brief_data and brief_data['table_of_authorities']:
        add_custom_heading(doc, 'Table of Authorities', level=1, font_size=18)
        for i, each in enumerate(brief_data['table_of_authorities'], start=1):
            add_custom_paragraph(doc, f"{i}. {each}", font_size=16)

    # Statement of Case
    if brief_data.get('statement_of_case'):
        add_custom_heading(doc, 'Statement of Case', level=1, font_size=18)
        add_custom_paragraph(doc, brief_data['statement_of_case'], font_size=16)

    # Summary of Arguments
    if brief_data.get('summary_of_arguments'):
        add_custom_heading(doc, 'Summary of Arguments', level=1, font_size=18)
        add_custom_paragraph(doc, brief_data['summary_of_arguments'], font_size=16)

    # Arguments
    if 'brief_arguments' in brief_data and brief_data['brief_arguments']:
        add_custom_heading(doc, 'Arguments', level=1, font_size=18)
        for argument in brief_data['brief_arguments']:
            add_custom_heading(doc, argument['title'], level=2, font_size=17)
            add_custom_paragraph(doc, argument['description'], font_size=16)

    # Conclusion
    if brief_data.get('conclusion'):
        add_custom_heading(doc, 'Conclusion', level=1, font_size=18)
        add_custom_paragraph(doc, brief_data['conclusion'], font_size=16)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer



def add_custom_heading(doc, text, level, font_name="Times New Roman", font_size=12):
    """Adds a heading with custom font name and size, centered."""
    # Add a heading (which internally is a paragraph)
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_custom_paragraph(doc, text, font_name="Times New Roman", font_size=12):
    """Adds a paragraph with custom font name and size."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)


def generate_cover_page(doc, data):
    add_custom_heading(doc, "NO.2013-01", level=5, font_size=16)
    add_custom_heading(doc, "------------------------------------------------", level=3, font_size=16)
    add_custom_heading(doc, "IN THE", level=5, font_size=16)
    add_custom_heading(doc, "SUPREME COURT", level=0, font_size=26)
    court_term = data.get("court_term", "Court term") 
    add_custom_heading(doc, court_term, level=0, font_size=16)
    add_custom_heading(doc, "--------------", level=3, font_size=16)
    add_custom_heading(doc, data.get('petitioner_name', ''), level=3, font_size=16)
    add_custom_heading(doc, data.get('V.', ''), level=3, font_size=12)
    add_custom_heading(doc, data.get('respondent_name', ''), level=3, font_size=16)
    add_custom_heading(doc, "--------------", level=3, font_size=12)
    add_custom_heading(doc, data.get('title_of_brief', ''), level=3, font_size=16)
    add_custom_heading(doc, "--------------", level=3, font_size=12)
    add_custom_heading(doc, data.get('petitioner_name', ''), level=3, font_size=16)
    add_custom_heading(doc, "Attorneys for the case", level=3, font_size=18)
    add_custom_heading(doc, data.get('ATTORNEYS FOR THE RESPONDENT:', ''), level=3, font_size=16)
    
    # Attorneys
    attorneys = data.get('attorneys', [])
    if attorneys:
        attorneys_str = ', '.join(attorneys)
        add_custom_heading(doc, attorneys_str, level=3, font_size=12)

    return doc






    